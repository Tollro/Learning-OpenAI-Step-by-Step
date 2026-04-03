from docx import Document
import os

def extract_text_from_doc_table(doc_path):
    # 新增 提取表格内容
    text = ""
    try:
        doc = Document(doc_path)
        # 提取段落
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text += para_text + "\n"
        # 提取表格（可选）
        for table in doc.tables:
            for row in table.rows:
                row_text = ""
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text += cell_text + "\t"
                if row_text.strip():
                    text += row_text.strip() + "\n"
    except Exception as e:
        print(f"Error extracting text from {doc_path}: {e}")
        return ""
    return text


def extract_text_from_doc(doc_path):
    """
    Extract text from a DOCX file, skipping blank or whitespace-only paragraphs.

    Args:
        doc_path (str): Path to the DOCX file.

    Returns:
        str: Extracted text from all paragraphs.
    """
    text = ""
    try:
        doc = Document(doc_path)
        for para in doc.paragraphs:
            para_text = para.text.strip()  # 去除首尾空白字符
            if para_text:                  # 如果去除后不为空
                text += para_text + "\n"
    except Exception as e:
        print(f"Error extracting text from {doc_path}: {e}")
        return ""
    return text

def process_docx(doc_path):
    """
    Process a DOCX file: extract text and save to TXT.

    Args:
        doc_path (str): Path to the DOCX file.
    """
    if not os.path.exists(doc_path):
        print(f"File {doc_path} does not exist.")
        return

    extracted_text = extract_text_from_doc(doc_path)

    # Save to a text file with the same name
    txt_path = os.path.splitext(doc_path)[0] +"docu"+ ".txt"
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(extracted_text)

    print(f"Extracted text saved to {txt_path}")

if __name__ == "__main__":
    process_docx("./RAG/retrieval/txt_extract/人事管理流程.docx")
